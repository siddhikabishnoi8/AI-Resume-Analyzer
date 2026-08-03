import re
import os
import pdfplumber
import docx

def extract_text(filepath):
    """Extract text from PDF or DOCX file."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(filepath)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(filepath):
    text = []
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        # Check tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return "\n".join(text)

def parse_contact_info(text, nlp=None):
    """Extract candidate name, email, phone, and profile links."""
    # Email Regex
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    email = emails[0] if emails else None
    
    # Phone Regex (look for standard format: +1 (555) 019-2834 or +91 9876543210 or 555-019-2834)
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    
    # Re-verify and clean general number blocks, ignoring year ranges
    raw_phones = re.findall(r'\+?\d[\d -]{8,15}\d', text)
    clean_phones = []
    for p in raw_phones:
        p_clean = p.strip()
        # Skip if it is a year range like "2020 - 2022" or "2018-2020"
        if re.match(r'^\d{4}\s*[-–—]\s*\d{4}$', p_clean):
            continue
        # Ensure it has enough digits
        digits_only = re.sub(r'\D', '', p_clean)
        if len(digits_only) >= 7 and len(digits_only) <= 15:
            clean_phones.append(p_clean)
            
    phone = clean_phones[0] if clean_phones else (phones[0] if phones else None)
    
    # Social profiles
    github_pattern = r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+'
    github = re.search(github_pattern, text, re.IGNORECASE)
    github_url = github.group(0) if github else None
    
    linkedin_pattern = r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+'
    linkedin = re.search(linkedin_pattern, text, re.IGNORECASE)
    linkedin_url = linkedin.group(0) if linkedin else None
    
    # Name extraction using spaCy PERSON entity or heuristics
    name = "Unknown Candidate"
    if nlp:
        try:
            # We scan the first 1000 characters for name
            doc = nlp(text[:1000])
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    candidate_name = ent.text.strip()
                    # Clean and validate name: 2-3 words, no newlines
                    if len(candidate_name.split()) >= 2 and len(candidate_name.split()) <= 4 and '\n' not in candidate_name:
                        name = candidate_name
                        break
        except Exception as e:
            print(f"spaCy NER Name extraction failed: {e}")
            
    if name == "Unknown Candidate":
        # Fallback to heuristic: scan first 5 non-empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines[:5]:
            # Exclude lines containing contact details or typical headers
            if (len(line.split()) >= 2 and len(line.split()) <= 4 and 
                len(line) < 50 and 
                '@' not in line and 
                'resume' not in line.lower() and 
                'curriculum' not in line.lower() and
                not any(char.isdigit() for char in line) and
                not any(x in line for x in [':', '/', '\\', '|', ',', '.'])):
                name = line
                break
                
    return {
        'name': name,
        'email': email,
        'phone': phone,
        'github': github_url,
        'linkedin': linkedin_url
    }

def parse_education(text):
    """Extract education details like Degree, Major, University, and Year."""
    education_keywords = ['education', 'academic', 'qualification', 'degree', 'university', 'college', 'school']
    
    # Common degrees pattern
    degree_patterns = [
        r'(Bachelor\s+of\s+[a-zA-Z\s]+|B\.S\.|B\.Sc\.|B\.A\.|B\.Tech|B\.E\.|B\.B\.A\.)',
        r'(Master\s+of\s+[a-zA-Z\s]+|M\.S\.|M\.Sc\.|M\.A\.|M\.Tech|M\.E\.|M\.B\.A\.)',
        r'(Ph\.D\.|PhD|Doctor\s+of\s+Philosophy)',
        r'(Associate\s+of\s+[a-zA-Z\s]+|A\.S\.|A\.A\.)'
    ]
    
    found_degrees = []
    for pattern in degree_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if match.strip() not in found_degrees:
                found_degrees.append(match.strip())
                
    # Quick clean
    cleaned_degrees = []
    for deg in found_degrees:
        # Standardize representation
        deg_lower = deg.lower()
        if 'ph.d' in deg_lower or 'phd' in deg_lower:
            cleaned_degrees.append('PhD')
        elif 'master' in deg_lower or 'm.s' in deg_lower or 'm.tech' in deg_lower or 'mba' in deg_lower:
            cleaned_degrees.append('Master\'s Degree')
        elif 'bachelor' in deg_lower or 'b.s' in deg_lower or 'b.tech' in deg_lower or 'b.e' in deg_lower:
            cleaned_degrees.append('Bachelor\'s Degree')
        elif 'associate' in deg_lower or 'a.s' in deg_lower:
            cleaned_degrees.append('Associate\'s Degree')
            
    # Default fallback to "Bachelor's Degree" if words like "University" or "College" are present but no specific degree matched
    if not cleaned_degrees:
        if any(kw in text.lower() for kw in ['university', 'college', 'institute of technology']):
            cleaned_degrees.append('Bachelor\'s Degree')
            
    return list(set(cleaned_degrees)) if cleaned_degrees else ["High School Diploma"]

def parse_experience(text):
    """Estimate years of experience based on date ranges in text."""
    # Look for date patterns: e.g., 2018 - 2022, Jan 2019 - Present, 05/2020 to 08/2023
    months = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2})'
    years = r'(?:19\d{2}|20\d{2}|\d{2})'
    separator = r'(?:\s*[-–—to]\s*|\s+)'
    
    date_pattern = re.compile(
        rf'(?:{months}\s+)?({years}){separator}(?:(?:{months}\s+)?({years})|Present|Current|Now)',
        re.IGNORECASE
    )
    
    matches = date_pattern.findall(text)
    total_years = 0
    intervals = []
    
    current_year = datetime = 2026 # Use current local year from environment (2026)
    
    for start, end in matches:
        try:
            # Clean years
            s_yr = int(start) if len(start) == 4 else int("20" + start if int(start) < 50 else "19" + start)
            
            if not end or end.lower() in ['present', 'current', 'now']:
                e_yr = current_year
            else:
                e_yr = int(end) if len(end) == 4 else int("20" + end if int(end) < 50 else "19" + end)
            
            # Simple boundary checks
            if 1970 < s_yr <= current_year and 1970 < e_yr <= current_year and e_yr >= s_yr:
                diff = e_yr - s_yr
                # Check overlapping ranges heuristically: if we have separate jobs, add them up
                intervals.append((s_yr, e_yr))
        except ValueError:
            continue
            
    # Calculate non-overlapping union of intervals
    if intervals:
        intervals.sort(key=lambda x: x[0])
        merged = [intervals[0]]
        for current in intervals[1:]:
            prev_start, prev_end = merged[-1]
            curr_start, curr_end = current
            if curr_start <= prev_end:
                # Overlap, merge
                merged[-1] = (prev_start, max(prev_end, curr_end))
            else:
                # No overlap
                merged.append(current)
        total_years = sum(end - start for start, end in merged)
        
    # Heuristics: search for "X+ years", "X years of experience"
    exp_text_pattern = re.search(r'(\d+)\+?\s*years?\s+(?:of\s+)?experience', text, re.IGNORECASE)
    if exp_text_pattern:
        text_years = int(exp_text_pattern.group(1))
        total_years = max(total_years, text_years)
        
    # If no dates found, let's check for work sections. If they exist but no dates can be parsed, default to a minimum standard (e.g. 1-2 years) or 0.
    if total_years == 0 and any(kw in text.lower() for kw in ['experience', 'work history', 'professional background']):
        total_years = 1 # default minimum if work history sections exist but dates aren't parsable
        
    return total_years
